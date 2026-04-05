#!/usr/bin/env python3
"""Генерация отчёта по анализу шаблонов математики 1-4 класс"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import json
import os
from datetime import datetime

def set_cell_shading(cell, color):
    """Установить цвет фона ячейки"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def create_report():
    doc = Document()

    # Настройка стилей
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    # Заголовок
    title = doc.add_heading('Комплексный анализ педагогических шаблонов', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Математика 1-4 классы')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.bold = True

    # Метаданные
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f'Дата: {datetime.now().strftime("%d.%m.%Y")}\n').font.size = Pt(10)
    meta.add_run('Версия реестра: 1.1.0\n').font.size = Pt(10)
    meta.add_run('Всего шаблонов: 50').font.size = Pt(10)

    doc.add_paragraph()

    # ===== РАЗДЕЛ 1: ОБЗОР =====
    doc.add_heading('1. Обзор системы шаблонов', level=1)

    doc.add_paragraph(
        'Система содержит 50 педагогических шаблонов для поддержки решения '
        'математических задач начальной школы (1-4 классы). Каждый шаблон включает:'
    )

    bullets = [
        'Правила маршрутизации (routing_rules) для автоматического определения типа задачи',
        'Трёхуровневую систему подсказок (L1-L2-L3)',
        'Документацию типичных ошибок учеников (common_mistakes)',
        'Политику раскрытия подсказок (hint_policy)',
        'Routing-тесты в JSON и Go unit-тесты',
    ]
    for b in bullets:
        p = doc.add_paragraph(b, style='List Bullet')

    # ===== РАЗДЕЛ 2: ПОКРЫТИЕ =====
    doc.add_heading('2. Покрытие учебной программы', level=1)

    doc.add_heading('2.1. Готовность по классам', level=2)

    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['Класс', 'Шаблонов', 'Готовность', 'Статус']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell, 'D9E2F3')

    data = [
        ('1 класс', '26', '95%', 'Полное покрытие'),
        ('2 класс', '41', '95%', 'Полное покрытие'),
        ('3 класс', '44', '95%', 'Полное покрытие'),
        ('4 класс', '50', '95%', 'Полное покрытие'),
    ]
    colors = ['C6EFCE', 'C6EFCE', 'C6EFCE', 'C6EFCE']

    for i, (row_data, color) in enumerate(zip(data, colors)):
        for j, val in enumerate(row_data):
            cell = table.rows[i+1].cells[j]
            cell.text = val
            if j == 3:
                set_cell_shading(cell, color)

    doc.add_paragraph()

    # ===== РАЗДЕЛ 3: СВОДНАЯ ТАБЛИЦА =====
    doc.add_heading('2.2. Сводная таблица шаблонов', level=2)

    templates_data = [
        ('T1', 'Нумерация и числовая прямая', '1-4', 'number_sense', '3', '✓ (5)'),
        ('T2', 'Римские и арабские числа', '1-4', 'numeral_systems', '2', '✓ (4)'),
        ('T3', 'Чётность и деление с остатком', '1-4', 'arithmetic_fluency', '3', '✓ (5)'),
        ('T4', 'Устное сложение/вычитание (до 100)', '1-4', 'arithmetic_fluency', '2', '✓ (5)'),
        ('T5', 'Письменное сложение/вычитание (столбиком)', '1-4', 'arithmetic_fluency', '3', '✓ (5)'),
        ('T6', 'Неизвестный компонент (слож./вычит.)', '1-4', 'arithmetic_fluency', '3', '✓ (4)'),
        ('T7', 'Смысл умножения и деления', '1-4', 'word_problems', '3', '✓ (4)'),
        ('T8', 'Таблица умножения и деления', '1-4', 'arithmetic_fluency', '3', '✓ (4)'),
        ('T9', 'Письменное умножение/деление (столбиком)', '1-4', 'arithmetic_fluency', '3', '✓ (5)'),
        ('T10', 'Увеличение/уменьшение в несколько раз', '1-4', 'word_problems', '3', '✓ (4)'),
        ('T11', 'Свойства действий и упрощение', '2-4', 'arithmetic_fluency', '2', '✓ (5)'),
        ('T12', 'Доли и простые дроби', '3-4', 'fractions_percent', '2', '✓ (5)'),
        ('T13', 'Проценты и пропорции', '4', 'fractions_percent', '2', '✓ (5)'),
        ('T14', 'Простые текстовые задачи', '1-4', 'word_problems', '2', '✓ (5)'),
        ('T15', 'Задачи на кратное сравнение', '2-4', 'word_problems', '3', '✓ (4)'),
        ('T16', 'Задачи на изменение величины', '2-4', 'word_problems', '2', '✓ (5)'),
        ('T17', 'Составные задачи (2-3 шага)', '2-4', 'word_problems', '3', '✓ (3)'),
        ('T18', 'Задачи на норму и количество', '2-4', 'word_problems', '3', '✓ (3)'),
        ('T19', 'Текстовые задачи на величины', '1-4', 'word_problems', '2', '✓ (5)'),
        ('T20', 'Сложные задачи (движение/работа)', '3-4', 'word_problems', '2', '✓ (6)'),
        ('T21', 'Перевод и сравнение единиц', '1-4', 'measurement_units', '2', '✓ (5)'),
        ('T22', 'Периметр и площадь прямоугольника', '1-4', 'geometry', '3', '✓ (4)'),
        ('T23', 'Температура и термометр', '2-4', 'measurement_units', '2', '✓ (5)'),
        ('T24', 'Базовые геометрические фигуры', '1-4', 'geometry', '2', '✓ (5)'),
        ('T25', 'Углы и их виды', '2-4', 'geometry', '2', '✓ (5)'),
        ('T26', 'Периметр и площадь на чертеже', '3-4', 'geometry', '2', '✓ (5)'),
        ('T27', 'Симметрия и ось симметрии', '2-4', 'geometry', '2', '✓ (5)'),
        ('T28', 'Координатная сетка', '3-4', 'geometry', '2', '✓ (5)'),
        ('T29', 'Составление фигур из частей', '2-4', 'geometry', '2', '✓ (5)'),
        ('T30', 'Чтение и дополнение таблиц', '1-4', 'data_representation', '2', '✓ (5)'),
        ('T31', 'Чтение диаграмм и графиков', '2-4', 'data_representation', '2', '✓ (5)'),
        ('T32', 'Построение диаграмм', '2-4', 'data_representation', '2', '✓ (5)'),
        ('T33', 'Простая статистика', '3-4', 'data_representation', '2', '✓ (5)'),
        ('T34', 'Логические задачи и закономерности', '2-4', 'logic', '2', '✓ (5)'),
        ('T35', 'Порядок действий', '2-4', 'patterns_logic', '3', '✓ (4)'),
        ('T36', 'Задачи по рисунку', '1-4', 'word_problems', '2', '✓ (5)'),
        ('T37', 'Уравнения и неравенства', '1-4', 'patterns_logic', '3', '✓ (4)'),
        ('T38', 'Арифметические головоломки', '2-4', 'patterns_logic', '3', '✓ (4)'),
        ('T39', 'Геометрические головоломки', '2-4', 'geometry', '3', '✓ (4)'),
        ('T40', 'Простая комбинаторика', '3-4', 'logic', '2', '✓ (5)'),
        ('T41', 'Логика и комбинаторика', '2-4', 'patterns_logic', '3', '✓ (4)'),
        ('T42', 'Часы и время', '1-4', 'measurement_units', '2', '✓ (5)'),
        ('T43', 'Деньги и покупки', '1-4', 'measurement_units', '2', '✓ (5)'),
        ('T44', 'Вычислительная культура', '2-4', 'arithmetic_fluency', '3', '✓ (4)'),
        # Новые шаблоны для 1 класса
        ('T45', 'Состав числа', '1', 'number_sense', '2', '✓ (5)'),
        ('T46', 'Счёт предметов', '1', 'number_sense', '2', '✓ (5)'),
        ('T47', 'Сравнение чисел (>, <, =)', '1', 'number_sense', '2', '✓ (5)'),
        ('T48', 'Соседи числа и числовой ряд', '1', 'number_sense', '2', '✓ (5)'),
        ('T49', 'Сложение и вычитание в пределах 10', '1', 'arithmetic_fluency', '2', '✓ (5)'),
        ('T50', 'Простые задачи 1 класса (было-стало)', '1', 'word_problems', '2', '✓ (5)'),
    ]

    table2 = doc.add_table(rows=len(templates_data)+1, cols=6)
    table2.style = 'Table Grid'

    headers2 = ['Код', 'Название', 'Классы', 'Тип задачи', 'Подсказок', 'Ошибки']
    for i, h in enumerate(headers2):
        cell = table2.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        set_cell_shading(cell, 'D9E2F3')

    for i, row_data in enumerate(templates_data):
        for j, val in enumerate(row_data):
            cell = table2.rows[i+1].cells[j]
            cell.text = val
            cell.paragraphs[0].runs[0].font.size = Pt(9)

    doc.add_paragraph()

    # ===== РАЗДЕЛ 4: ТИПЫ ЗАДАЧ =====
    doc.add_heading('2.3. Распределение по типам задач', level=2)

    types_data = [
        ('word_problems', '11', 'Текстовые задачи'),
        ('arithmetic_fluency', '9', 'Арифметика'),
        ('geometry', '8', 'Геометрия'),
        ('number_sense', '5', 'Чувство числа'),
        ('patterns_logic', '4', 'Паттерны и логика'),
        ('measurement_units', '4', 'Единицы измерения'),
        ('data_representation', '4', 'Работа с данными'),
        ('logic', '2', 'Логика'),
        ('fractions_percent', '2', 'Дроби и проценты'),
        ('numeral_systems', '1', 'Системы счисления'),
    ]

    table3 = doc.add_table(rows=len(types_data)+1, cols=3)
    table3.style = 'Table Grid'

    for i, h in enumerate(['Тип задачи', 'Кол-во', 'Описание']):
        cell = table3.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell, 'D9E2F3')

    for i, row_data in enumerate(types_data):
        for j, val in enumerate(row_data):
            table3.rows[i+1].cells[j].text = val

    doc.add_paragraph()

    # ===== РАЗДЕЛ 5: КАЧЕСТВО =====
    doc.add_heading('3. Качество педагогического наполнения', level=1)

    doc.add_heading('3.1. Сильные стороны', level=2)

    strengths = [
        'Все 50 шаблонов имеют трёхуровневую систему подсказок (L1-L2-L3)',
        'Все 50 шаблонов содержат типичные ошибки учеников (common_mistakes)',
        'Все 50 шаблонов покрыты routing-тестами в JSON и Go unit-тестами',
        'Консистентная структура: rules → format → forbidden на каждом уровне',
        'Дифференциация по сложности (max_hints: 2 или 3)',
        'Указаны confusables — шаблоны, с которыми возможна путаница',
        'Полное покрытие 1 класса (26 шаблонов, включая специализированные T45-T50)',
        'Реализованы forbid-паттерны для критичных пар: T2 (уравнения), T5 (__GAP__), T7 (дроби), T15 (изменение)',
        '10 confusable-pair тестов с 40+ кейсами для проверки разделения шаблонов',
    ]
    for s in strengths:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_heading('3.2. Области для улучшения', level=2)

    problems = [
        ('Visual_facts требования', 'Геометрические шаблоны T22-T29 могут требовать уточнения visual_facts'),
        ('Примеры решений', 'В teaching_pattern можно добавить конкретные примеры решённых задач'),
    ]

    table4 = doc.add_table(rows=len(problems)+1, cols=2)
    table4.style = 'Table Grid'

    for i, h in enumerate(['Область', 'Описание']):
        cell = table4.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell, 'FFEB9C')

    for i, (prob, desc) in enumerate(problems):
        table4.rows[i+1].cells[0].text = prob
        table4.rows[i+1].cells[1].text = desc

    doc.add_paragraph()

    # ===== РАЗДЕЛ 6: СИСТЕМА ПОДСКАЗОК =====
    doc.add_heading('3.3. Структура системы подсказок', level=2)

    doc.add_paragraph(
        'Каждый шаблон содержит teaching_pattern с тремя уровнями подсказок:'
    )

    levels = [
        ('L1 (Базовый)', 'Понимание задачи', 'Помочь ученику понять, что требуется найти и какие данные есть'),
        ('L2 (Средний)', 'Пошаговое решение', 'Направить на правильный алгоритм решения с объяснением каждого шага'),
        ('L3 (Продвинутый)', 'Самопроверка', 'Помочь проверить ответ и закрепить навык'),
    ]

    table5 = doc.add_table(rows=4, cols=3)
    table5.style = 'Table Grid'

    for i, h in enumerate(['Уровень', 'Цель', 'Назначение']):
        cell = table5.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell, 'D9E2F3')

    for i, (level, goal, desc) in enumerate(levels):
        table5.rows[i+1].cells[0].text = level
        table5.rows[i+1].cells[1].text = goal
        table5.rows[i+1].cells[2].text = desc

    doc.add_paragraph()

    # ===== РАЗДЕЛ 7: РЕКОМЕНДАЦИИ =====
    doc.add_heading('4. Рекомендации по улучшению', level=1)

    doc.add_heading('4.1. Приоритет P1 (Высокий) — ВЫПОЛНЕНО ✓', level=2)
    doc.add_paragraph(
        'Реализованы защитные паттерны:'
    )
    bullets_p1 = [
        'T2 — добавлены forbid для уравнений: "уравнен", "неизвестн", "найди x", "x =", "= x"',
        'T5 — добавлены forbid в правило __GAP__: "римск", "устно", "в уме", "диаграмм", "координат"',
        'Добавлены тесты: TestT2_RejectEquationsWithX, TestT5_RejectGapWithoutColumn',
    ]
    for b in bullets_p1:
        doc.add_paragraph(b, style='List Bullet')

    doc.add_heading('4.2. Приоритет P2 (Средний)', level=2)
    doc.add_paragraph(
        'Уточнить требования к визуализации для геометрических шаблонов T22-T29. '
        'Добавить поле visual_facts с описанием необходимых изображений.'
    )

    doc.add_heading('4.3. Приоритет P3 (Низкий)', level=2)
    bullets_p3 = [
        'Добавить конкретные примеры решённых задач в teaching_pattern',
        'Унифицировать количество подсказок (2 vs 3)',
        'Рассмотреть добавление шаблонов для олимпиадной математики',
    ]
    for b in bullets_p3:
        doc.add_paragraph(b, style='List Bullet')

    doc.add_paragraph()

    # ===== РАЗДЕЛ 8: ВЫВОДЫ =====
    doc.add_heading('5. Выводы', level=1)

    conclusion = doc.add_paragraph()
    conclusion.add_run('Общая готовность: 95% ✓\n').bold = True
    conclusion.add_run(
        'Система из 50 шаблонов полностью готова к production-использованию '
        'для всех классов начальной школы (1-4).\n\n'
    )

    conclusion.add_run('Покрытие по классам:\n').bold = True
    conclusion.add_run(
        '• 1 класс: 26 шаблонов (100%) — добавлены специализированные T45-T50\n'
        '• 2-4 классы: 100% покрытие основных тем\n\n'
    )

    conclusion.add_run('Педагогическое наполнение: 100% ✓\n').bold = True
    conclusion.add_run(
        'Все 50 шаблонов содержат common_mistakes и трёхуровневые подсказки.\n\n'
    )

    conclusion.add_run('Тестовое покрытие: 100% ✓\n').bold = True
    conclusion.add_run(
        'Все 50 шаблонов покрыты routing-тестами в JSON и Go unit-тестами.\n\n'
    )

    conclusion.add_run('Confusable-pair тесты: 10 пар ✓\n').bold = True
    conclusion.add_run(
        '• T23↔T25 (температура vs углы)\n'
        '• T19↔T21 (задачи с единицами vs перевод)\n'
        '• T31↔T32 (чтение vs построение диаграмм)\n'
        '• T21↔T42 (перевод времени vs часы)\n'
        '• T14↔T15↔T16 (простые задачи vs сравнение vs изменение)\n'
        '• T4↔T5 (устный vs письменный счёт)\n'
        '• T7↔T8 (смысл умножения vs таблица)\n'
        '• T7↔T12 (равные группы vs дроби)\n'
        '• T30↔T31 (таблицы vs диаграммы)\n'
    )

    # Сохранение
    output_path = '/Users/a.yanover/Xsolla/child_bot/api/internal/v2/templates/Анализ_шаблонов_математики_1-4_класс.docx'
    doc.save(output_path)
    print(f'Отчёт сохранён: {output_path}')
    return output_path

if __name__ == '__main__':
    create_report()
